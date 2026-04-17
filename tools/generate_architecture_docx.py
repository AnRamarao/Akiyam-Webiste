#!/usr/bin/env python3
"""
Generate AIKYAM Architecture Document (Word .docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

GOLD = RGBColor(0xFF, 0xD7, 0x00)
DARK_BG = RGBColor(0x1A, 0x1A, 0x2E)
DARK_BLUE = RGBColor(0x1A, 0x36, 0x5D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A365D')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def main():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = DARK_BLUE

    # ===== TITLE PAGE =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AIKYAM')
    run.font.size = Pt(42)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Architecture Document')
    run.font.size = Pt(20)
    run.font.color.rgb = MED_GRAY

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ('Version: ', '1.0'),
        ('  |  Date: ', 'April 17, 2026'),
        ('  |  Author: ', 'Ramarao / Engineering'),
    ]:
        run = meta.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = MED_GRAY
        run = meta.add_run(value)
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = DARK_BLUE

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Overview',
        '3. Technology Stack',
        '4. File & Directory Structure',
        '5. Frontend Architecture',
        '6. Authentication & Authorization',
        '7. Event Management System',
        '8. Database Schema',
        '9. Security Architecture',
        '10. Data Flow',
        '11. Build & Deployment',
        '12. Environment Setup Guide',
        '13. Role Access Matrix',
        '14. Performance & Accessibility',
        '15. Future Considerations',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'AIKYAM is a community website for a 501(c)(3) non-profit organization based in California. '
        'The platform manages community events, team information, vendor directories, photo galleries, '
        'and member engagement \u2014 all built on a zero-cost infrastructure using static HTML/CSS/JS '
        'hosted on GoDaddy with Supabase (free tier) as the backend.'
    )

    doc.add_heading('Key Design Decisions', level=2)
    add_styled_table(doc,
        ['Decision', 'Rationale'],
        [
            ['Static site (no framework)', 'Zero hosting cost on GoDaddy, fast load times, simple maintenance'],
            ['Supabase free tier', 'PostgreSQL + Auth + RLS \u2014 no backend server needed'],
            ['Client-side only', 'No server runtime required \u2014 all logic runs in browser'],
            ['IIFE module pattern', 'No build toolchain needed, explicit dependency management'],
            ['JSON fallback', 'Graceful degradation when Supabase is unreachable'],
        ],
        col_widths=[5, 12]
    )

    doc.add_page_break()

    # ===== 2. SYSTEM OVERVIEW =====
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph(
        'The system operates entirely client-side. The browser loads static HTML/CSS/JS from GoDaddy, '
        'then JavaScript modules communicate with Supabase for authentication and database operations. '
        'All authorization is enforced at two levels: client-side (for UX) and server-side via '
        'Row-Level Security policies (for actual security).'
    )

    doc.add_heading('Architecture Diagram', level=2)
    add_code_block(doc, """
+----------------------------------------------------------+
|                     USER'S BROWSER                        |
|                                                           |
|  +-------------+  +----------+  +-----------+             |
|  | index.html  |  | auth.js  |  | events.js |             |
|  | script.js   |  | AIKYAM_  |  | AIKYAM_   |             |
|  | styles.css  |  | Auth     |  | Events    |             |
|  +------+------+  +----+-----+  +-----+-----+             |
|         |              |              |                    |
+---------+--------------+--------------+-------------------+
          |              |              |
   +--------------+         +--------------+
   | GoDaddy      |         | Supabase     |
   | Static Host  |         | (Free Tier)  |
   |              |         |              |
   | HTML, CSS,   |         | Auth (GoTrue)|
   | JS, Images,  |         | PostgreSQL   |
   | JSON data    |         | RLS Policies |
   +--------------+         +--------------+
""")

    doc.add_heading('Module Dependency Graph', level=2)
    add_code_block(doc, """
                    config.js
                       |
                       v
  supabase-js (CDN) --> auth.js
                          |
                          | exposes getClient()
                          v
                     events.js
                          |
              +-----------+-----------+
              |                       |
              v                       v
         script.js             admin-events.js
    (public site rendering)   (admin dashboard)
""")

    doc.add_heading('Script Load Order (all defer)', level=2)
    add_styled_table(doc,
        ['Order', 'File', 'Purpose'],
        [
            ['1', 'supabase-js (CDN)', 'Supabase client library'],
            ['2', 'config.js', 'Credentials (SUPABASE_URL, SUPABASE_ANON_KEY)'],
            ['3', 'auth.js', 'Auth module, creates Supabase client'],
            ['4', 'events.js', 'Events module, shares auth client'],
            ['5', 'admin-events.js', 'Admin page logic (admin-events.html only)'],
            ['6', 'script.js', 'Main app initialization'],
        ],
        col_widths=[1.5, 5, 10]
    )

    doc.add_page_break()

    # ===== 3. TECHNOLOGY STACK =====
    doc.add_heading('3. Technology Stack', level=1)
    add_styled_table(doc,
        ['Layer', 'Technology', 'Details'],
        [
            ['Markup', 'HTML5', 'Semantic elements, ARIA attributes'],
            ['Styling', 'CSS3', 'Custom properties, Grid, Flexbox, BEM naming'],
            ['Logic', 'Vanilla JavaScript', 'ES2017+ (async/await), IIFE modules'],
            ['Fonts', 'Google Fonts', 'Montserrat (400, 700, 900)'],
            ['Auth', 'Supabase Auth (GoTrue)', 'Client-side SDK v2 via CDN'],
            ['Database', 'Supabase PostgreSQL', 'Free tier, Row-Level Security'],
            ['Hosting', 'GoDaddy', 'Static file hosting'],
            ['Images', 'WebP + JPEG/PNG', 'Responsive <picture> elements'],
            ['Build', 'Make + Python3', 'Image optimization, CSS/JS minification'],
            ['Version Control', 'Git / GitHub', 'Feature branch workflow'],
        ],
        col_widths=[3, 5, 9]
    )

    doc.add_heading('External Dependencies (CDN Only)', level=2)
    add_styled_table(doc,
        ['Package', 'CDN', 'Purpose'],
        [
            ['@supabase/supabase-js@2', 'jsdelivr.net', 'Database & auth client'],
            ['Montserrat font', 'fonts.googleapis.com', 'Typography'],
        ],
        col_widths=[5, 5, 7]
    )

    doc.add_page_break()

    # ===== 4. FILE STRUCTURE =====
    doc.add_heading('4. File & Directory Structure', level=1)
    add_code_block(doc, """
Akiyam-Webiste/
|-- index.html                  Homepage (hero, team, events, vendors)
|-- login.html                  Sign-in page
|-- signup.html                 Registration page
|-- reset-password.html         Password recovery
|-- admin-events.html           Event management dashboard
|-- gallery.html                Photo gallery
|-- vendors.html                Vendor directory
|-- contact.html                Contact form
|-- donate.html / join.html     Donations / Membership
|-- fillout.html / launch.html  Event registration / Special landing
|
|-- script.js                   Main app logic (~53KB)
|-- auth.js                     Auth module, AIKYAM_Auth (~19KB)
|-- events.js                   Events module, AIKYAM_Events (~8KB)
|-- admin-events.js             Admin events page logic (~10KB)
|-- config.js                   Supabase credentials (.gitignored)
|-- config.example.js           Credential template (committed)
|-- styles.css                  All styles (~82KB)
|
|-- common/
|   |-- header.html             Shared nav header
|   |-- footer.html             Shared footer
|
|-- data/
|   |-- upcomingEvents.json     Upcoming events (JSON fallback)
|   |-- completedEvents.json    Past events (JSON fallback)
|   |-- coreTeam.json           Executive committee
|   |-- boardMembers.json       Board of directors
|   |-- vendors.json            Vendor directory
|   |-- galleryImages.json      Gallery image manifest
|
|-- assets/
|   |-- branding/               Logo, backgrounds (PNG, WebP)
|   |-- events/                 Event photos & videos
|   |-- gallery/migrated/       Gallery photos (optimized WebP)
|   |-- people/                 Team/board member photos
|
|-- tools/                      Build scripts
|-- supabase-schema.sql         Database schema (idempotent)
|-- supabase-seed.sql           Test users & roles (.gitignored)
|-- Makefile                    Build automation
""")

    doc.add_page_break()

    # ===== 5. FRONTEND ARCHITECTURE =====
    doc.add_heading('5. Frontend Architecture', level=1)

    doc.add_heading('JavaScript Module Pattern (IIFE)', level=2)
    doc.add_paragraph(
        'All modules use the Immediately Invoked Function Expression pattern to avoid '
        'global namespace pollution. Each module exposes its public API on the window object.'
    )
    add_code_block(doc, """
(function () {
  'use strict';
  // Private state and functions
  var privateVar = 'hidden';
  function privateHelper() { /* ... */ }

  // Public API
  window.AIKYAM_ModuleName = {
    publicMethod: publicMethod
  };
})();
""")

    doc.add_heading('AIKYAM_Auth API', level=2)
    add_styled_table(doc,
        ['Method', 'Description'],
        [
            ['init()', 'Initialize auth, handle email confirmations'],
            ['getClient()', 'Return shared Supabase client instance'],
            ['getSession()', 'Get current auth session (or null)'],
            ['getProfile()', 'Get user profile with role'],
            ['signIn(email, password)', 'Sign in with credentials'],
            ['signUp(email, password, name)', 'Create new account'],
            ['signOut()', 'End session'],
            ['resetPassword(email)', 'Send password reset email'],
            ['updatePassword(newPassword)', 'Update password (from reset link)'],
            ['updateHeaderAuthState()', 'Refresh header login/user UI'],
            ['requireAuth(role)', 'Guard page by minimum role'],
            ['hasMinRole(userRole, required)', 'Check role >= required'],
            ['isAdmin(role)', 'Check GlobalAdmin or OrgAdmin'],
            ['isStaff(role)', 'Check admin or staff role'],
        ],
        col_widths=[6, 11]
    )

    doc.add_heading('AIKYAM_Events API', level=2)
    add_styled_table(doc,
        ['Method', 'Description'],
        [
            ['fetchPublicEvents()', 'Returns { upcoming: [], past: [] }'],
            ['fetchAllEvents()', 'All events in all statuses (admin use)'],
            ['createEvent(data)', 'Insert new draft event'],
            ['updateEvent(id, changes)', 'Update event fields'],
            ['publishEvent(id)', 'Draft \u2192 Published'],
            ['deactivateEvent(id)', 'Published \u2192 Deactivated'],
            ['deleteEvent(id)', 'Permanent delete'],
            ['schedulePublish(id, datetime)', 'Set future publish time'],
        ],
        col_widths=[6, 11]
    )

    doc.add_heading('CSS Architecture', level=2)
    doc.add_paragraph(
        'Methodology: BEM (Block__Element--Modifier). '
        'Theming: CSS custom properties with dark (default) / light modes. '
        'Responsive: Mobile-first with breakpoints at 768px and 1024px.'
    )

    add_styled_table(doc,
        ['Variable', 'Purpose', 'Dark Value', 'Light Value'],
        [
            ['--bg', 'Background', 'hsl(220,15%,8%)', 'hsl(0,0%,99%)'],
            ['--text', 'Primary text', 'hsl(220,20%,90%)', 'hsl(0,0%,15%)'],
            ['--accent', 'Gold accent', '#ffd700', '#ffd700'],
            ['--card', 'Card background', 'rgba(40,45,60,0.6)', 'rgba(255,255,255,0.6)'],
            ['--brd', 'Border color', 'hsl(220,15%,20%)', 'hsl(0,0%,80%)'],
        ],
        col_widths=[3, 4, 5, 5]
    )

    doc.add_page_break()

    # ===== 6. AUTHENTICATION =====
    doc.add_heading('6. Authentication & Authorization', level=1)

    doc.add_heading('Role Hierarchy (9 Levels)', level=2)
    add_styled_table(doc,
        ['Level', 'Role', 'Description'],
        [
            ['8', 'GlobalAdmin', 'Full system access, manage all users and roles'],
            ['7', 'OrgAdmin', 'Org-level admin, manage team, events, content, users'],
            ['6', 'ContentManager', 'Content lifecycle (publish, deactivate)'],
            ['5', 'ProgramManager', 'Event creation and management'],
            ['4', 'FinanceStaff', 'Financial access and reporting'],
            ['3', 'Volunteer', 'Event coordination, member-area access'],
            ['2', 'Donor', 'Donation history, donor-specific content'],
            ['1', 'Support', 'View/manage contact submissions'],
            ['0', 'Member', 'Default role, basic authenticated access'],
        ],
        col_widths=[1.5, 4.5, 11]
    )

    doc.add_heading('Sign-Up Flow', level=2)
    add_code_block(doc, """
User fills signup form
    |
    v
Rate limit check (5 attempts / 60 seconds)
    |
    v
supabase.auth.signUp({ email, password, data: { full_name } })
    |
    v
Supabase sends confirmation email
    |
    v
User clicks email link --> redirects to login.html#access_token=...&type=signup
    |
    v
auth.js detects hash fragment --> calls setSession() --> shows "Email confirmed!"
    |
    v
handle_new_user() trigger --> INSERT INTO profiles (role = 'Member')
""")

    doc.add_heading('Login Flow', level=2)
    add_code_block(doc, """
Email + Password --> Supabase validates --> Session created
                                                |
                                        Header updates:
                                        Login btn --> Avatar dropdown (name + role badge)
                                                |
                                        Role-gated UI elements appear/hide
""")

    doc.add_heading('Route Protection', level=2)
    doc.add_paragraph(
        'Protected pages call requireAuth(role) on load. '
        'No session \u2192 redirect to login.html. '
        'Session but insufficient role \u2192 redirect to index.html. '
        'Session + sufficient role \u2192 page loads normally.'
    )

    doc.add_heading('Conditional UI Visibility', level=2)
    doc.add_paragraph(
        'HTML elements use data-auth-required attribute to show/hide based on auth state. '
        'Supports values: "authenticated", "admin", "staff", or any specific role name.'
    )

    doc.add_page_break()

    # ===== 7. EVENT MANAGEMENT =====
    doc.add_heading('7. Event Management System', level=1)

    doc.add_heading('Event Lifecycle State Machine', level=2)
    add_code_block(doc, """
                         +===========+
          Create ------->||  DRAFT  ||
                         +====+======+
                              |
              +---------------+----------------+
              |                                |
        Manual Publish               Scheduled publish_at
        (ContentManager+)            (auto_publish_events)
              |                                |
              +---------------+----------------+
                              |
                         +====v======+
                         || PUBLISHED||<------+
                         +====+======+        |
                              |               |
              +---------------+--------+      |
              |                        |      |
         end_at passed          Manual deactivate
     (auto_expire_events)      (ContentManager+)
              |                        |      |
         +====v======+          +=====v=======+
         || EXPIRED  ||         || DEACTIVATED||
         || (visible ||         || (hidden)   ||---> Re-publish
         || as past) ||         +=============+
         +===========+
""")

    doc.add_heading('Event Data Model', level=2)
    add_styled_table(doc,
        ['Column', 'Type', 'Description'],
        [
            ['id', 'UUID (PK)', 'Auto-generated'],
            ['title', 'TEXT', 'Event name (required)'],
            ['description', 'TEXT', 'Full description'],
            ['location', 'TEXT', 'Venue (default: TBD)'],
            ['price', 'NUMERIC(10,2)', 'Ticket price (default: 0)'],
            ['img', 'TEXT', 'Image path'],
            ['tbd', 'BOOLEAN', 'Hide exact dates from public'],
            ['start_at', 'TIMESTAMPTZ', 'Event start'],
            ['end_at', 'TIMESTAMPTZ', 'Event end'],
            ['status', 'event_status', 'draft | published | deactivated | expired'],
            ['publish_at', 'TIMESTAMPTZ', 'Scheduled publish time'],
            ['expired_at', 'TIMESTAMPTZ', 'When auto-expired'],
            ['summary', 'TEXT', 'Shown on past event cards'],
            ['created_by', 'UUID (FK)', "Creator's profile ID"],
            ['updated_by', 'UUID (FK)', "Last editor's profile ID"],
            ['created_at', 'TIMESTAMPTZ', 'Row creation time'],
            ['updated_at', 'TIMESTAMPTZ', 'Last modification time'],
        ],
        col_widths=[3, 4, 10]
    )

    doc.add_heading('Admin Capabilities by Role', level=2)
    add_styled_table(doc,
        ['Role', 'View All', 'Create/Edit', 'Publish/Deactivate', 'Delete'],
        [
            ['GlobalAdmin', 'Yes', 'Yes', 'Yes', 'Yes'],
            ['OrgAdmin', 'Yes', 'Yes', 'Yes', 'Yes'],
            ['ContentManager', 'Yes', 'Yes', 'Yes', 'No'],
            ['ProgramManager', 'Yes', 'Own only', 'No', 'No'],
            ['Below PM', 'No', 'No', 'No', 'No'],
        ],
        col_widths=[4, 2.5, 3, 4, 2.5]
    )

    doc.add_heading('Data Fallback Strategy', level=2)
    doc.add_paragraph(
        'The site never shows a blank page. If Supabase is unreachable, events load from '
        'static JSON files (data/upcomingEvents.json and data/completedEvents.json). '
        'The format mappers ensure backward compatibility \u2014 rendering code works identically '
        'regardless of data source.'
    )

    doc.add_page_break()

    # ===== 8. DATABASE SCHEMA =====
    doc.add_heading('8. Database Schema', level=1)

    doc.add_heading('Entity Relationship', level=2)
    add_code_block(doc, """
+------------------+          +------------------+
|   auth.users     |    1:N   |  auth.identities |
|  (Supabase)      +----------+  (Supabase)      |
| id (PK)          |          | user_id (FK)      |
| email            |          | provider           |
| encrypted_pass   |          +------------------+
+--------+---------+
         | 1:1
+--------v---------+
|    profiles       |          1:N created_by
| id (PK, FK)      +--------------------------+
| email             |          1:N updated_by  |
| full_name         +---------------------+   |
| role (app_role)   |                     |   |
+-------------------+               +-----v---v---+
                                    |    events    |
                                    | id, title,   |
                                    | status, ...  |
                                    +--------------+
""")

    doc.add_heading('Profiles RLS Policies', level=2)
    add_styled_table(doc,
        ['Policy', 'Operation', 'Rule'],
        [
            ['Users read own', 'SELECT', 'auth.uid() = id'],
            ['Admins read all', 'SELECT', 'user_is_admin()'],
            ['Staff read all', 'SELECT', 'user_is_staff()'],
            ['Users update own', 'UPDATE', 'auth.uid() = id AND role unchanged'],
            ['GlobalAdmin update any', 'UPDATE', "user_has_role('GlobalAdmin')"],
            ['OrgAdmin update non-global', 'UPDATE', "user_has_role('OrgAdmin') AND target != GlobalAdmin"],
            ['GlobalAdmin delete', 'DELETE', "user_has_role('GlobalAdmin')"],
        ],
        col_widths=[5, 3, 9]
    )

    doc.add_heading('Events RLS Policies', level=2)
    add_styled_table(doc,
        ['Policy', 'Operation', 'Rule'],
        [
            ['Public read published', 'SELECT', "status = 'published'"],
            ['Public read expired', 'SELECT', "status = 'expired'"],
            ['Managers read all', 'SELECT', "user_has_min_role('ProgramManager')"],
            ['Managers create', 'INSERT', "user_has_min_role('ProgramManager')"],
            ['ContentManagers update any', 'UPDATE', "user_has_min_role('ContentManager')"],
            ['ProgramManagers update own', 'UPDATE', "user_has_min_role('PM') AND created_by = uid"],
            ['Admins delete', 'DELETE', 'user_is_admin()'],
        ],
        col_widths=[5, 3, 9]
    )

    doc.add_heading('Database Functions', level=2)
    add_styled_table(doc,
        ['Function', 'Type', 'Purpose'],
        [
            ['handle_new_user()', 'Trigger', 'Auto-create profile on auth signup'],
            ['update_updated_at()', 'Trigger', 'Auto-set updated_at on row changes'],
            ['user_has_role(role)', 'Helper', 'Exact role match check'],
            ['user_is_admin()', 'Helper', 'GlobalAdmin or OrgAdmin check'],
            ['user_is_staff()', 'Helper', 'Admin or manager/staff check'],
            ['user_has_min_role(role)', 'Helper', 'Hierarchy-aware role comparison'],
            ['auto_expire_events()', 'Lifecycle', 'Move past-end published to expired'],
            ['auto_publish_events()', 'Lifecycle', 'Move scheduled drafts to published'],
        ],
        col_widths=[5, 3, 9]
    )

    doc.add_page_break()

    # ===== 9. SECURITY =====
    doc.add_heading('9. Security Architecture', level=1)
    doc.add_paragraph('The system employs five layers of defense:')

    add_styled_table(doc,
        ['Layer', 'Mechanism', 'What It Protects Against'],
        [
            ['1', 'Content Security Policy (CSP)', 'XSS via injected scripts/styles'],
            ['2', 'Row-Level Security (RLS)', 'Unauthorized database access (server-side)'],
            ['3', 'Client-Side Auth Guards', 'Unauthorized page access (UX layer)'],
            ['4', 'Rate Limiting', 'Brute-force credential attacks'],
            ['5', 'Credential Isolation', 'Secrets in version control'],
        ],
        col_widths=[1.5, 5.5, 10]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Key insight: Client-side checks are for UX. Server-side RLS is for security. '
        'Even if someone opens browser DevTools, the database will reject unauthorized queries.'
    )
    run.bold = True

    doc.add_heading('Content Security Policy', level=2)
    add_code_block(doc, """
default-src 'self';
script-src  'self' https://cdn.jsdelivr.net;
style-src   'self' https://fonts.googleapis.com 'unsafe-inline';
font-src    https://fonts.gstatic.com;
img-src     'self' data:;
connect-src 'self' https://*.supabase.co;
""")

    doc.add_page_break()

    # ===== 10. DATA FLOW =====
    doc.add_heading('10. Data Flow', level=1)

    doc.add_heading('Public Page Load (Homepage)', level=2)
    add_code_block(doc, """
Browser                     GoDaddy              Supabase
   |--- GET index.html ------->|                     |
   |<-- HTML + scripts --------|                     |
   |--- GET *.js, *.css ------>|                     |
   |<-- static files ----------|                     |
   |  [DOMContentLoaded]       |                     |
   |--- GET common/header.html>|                     |
   |--- auth.getSession() ---------------------------->|
   |<-- session (or null) -----------------------------|
   |--- fetchPublicEvents() -------------------------->|
   |<-- { upcoming: [...], past: [...] } --------------|
   |--- GET data/*.json ------>|                     |
   |<-- team, vendors, etc ----|                     |
   |  [Render page]            |                     |
""")

    doc.add_heading('Admin Event Creation', level=2)
    add_code_block(doc, """
Admin Browser                                  Supabase
   |--- requireAuth('ProgramManager') ----------->|
   |<-- session + profile (role verified) ---------|
   |  [User fills form, clicks Save]              |
   |--- createEvent({ title, desc, ... }) ------->|
   |   [RLS: user_has_min_role('ProgramManager')] |
   |<-- { data: event_row } ----------------------|
   |--- fetchAllEvents() ----------------------->|
   |<-- all events -------------------------------|
   |  [Refresh list]                              |
""")

    doc.add_page_break()

    # ===== 11. BUILD & DEPLOYMENT =====
    doc.add_heading('11. Build & Deployment', level=1)

    doc.add_heading('Makefile Targets', level=2)
    add_styled_table(doc,
        ['Command', 'Purpose'],
        [
            ['make serve', 'Start dev server on port 8000'],
            ['make optimize', 'Convert images to WebP, generate responsive variants'],
            ['make minify', 'Minify CSS + all JS files'],
            ['make build', 'Full production build (optimize + minify)'],
            ['make clean', 'Remove all generated files'],
        ],
        col_widths=[4, 13]
    )

    doc.add_heading('Build Artifacts', level=2)
    doc.add_paragraph(
        'styles.min.css, script.min.js, auth.min.js, events.min.js, admin-events.min.js, '
        'WebP image variants (800w, 1600w), optimized people photos.'
    )

    doc.add_heading('Deployment Checklist', level=2)
    checklist = [
        'Run make build for production assets',
        'Verify config.js is NOT in git (.gitignored)',
        'Verify supabase-seed.sql is NOT in git',
        'Test all auth flows (signup, login, reset, logout)',
        'Test event CRUD with each role level',
        'Test JSON fallback (block Supabase URL)',
        'Test mobile responsive at 375px, 768px, 1440px',
        'Test dark/light theme toggle',
        'Check browser console for errors',
        'Upload to GoDaddy hosting',
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ===== 12. ENVIRONMENT SETUP =====
    doc.add_heading('12. Environment Setup Guide', level=1)

    doc.add_heading('For New Developers', level=2)
    steps = [
        'Clone the repository: git clone <repo-url> && cd Akiyam-Webiste',
        'Create config.js: cp config.example.js config.js (edit with Supabase credentials)',
        'Start dev server: make serve (opens at http://localhost:8000)',
        'Supabase setup: Run supabase-schema.sql in SQL Editor, then supabase-seed.sql',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Test Accounts', level=2)
    add_styled_table(doc,
        ['Email', 'Role', 'Password'],
        [
            ['rjramarao@gmail.com', 'GlobalAdmin', 'Menifee@92585'],
            ['orgadmin@aikyamusa.org', 'OrgAdmin', 'Menifee@92585'],
            ['contentmanager@aikyamusa.org', 'ContentManager', 'Menifee@92585'],
            ['programmanager@aikyamusa.org', 'ProgramManager', 'Menifee@92585'],
            ['financestaff@aikyamusa.org', 'FinanceStaff', 'Menifee@92585'],
            ['volunteer@aikyamusa.org', 'Volunteer', 'Menifee@92585'],
            ['donor@aikyamusa.org', 'Donor', 'Menifee@92585'],
            ['support@aikyamusa.org', 'Support', 'Menifee@92585'],
            ['member@aikyamusa.org', 'Member', 'Menifee@92585'],
        ],
        col_widths=[7, 4, 4]
    )

    doc.add_page_break()

    # ===== 13. ROLE ACCESS MATRIX =====
    doc.add_heading('13. Role Access Matrix', level=1)

    roles = ['GlbAdm', 'OrgAdm', 'CntMgr', 'PrgMgr', 'FinStf', 'Volntr', 'Donor', 'Supprt', 'Member']
    features = [
        ['View public site',        'Y','Y','Y','Y','Y','Y','Y','Y','Y'],
        ['Login / Session',         'Y','Y','Y','Y','Y','Y','Y','Y','Y'],
        ['View own profile',        'Y','Y','Y','Y','Y','Y','Y','Y','Y'],
        ['Update own name',         'Y','Y','Y','Y','Y','Y','Y','Y','Y'],
        ['See Admin nav link',      'Y','Y','Y','Y','-','-','-','-','-'],
        ['View all events (admin)', 'Y','Y','Y','Y','-','-','-','-','-'],
        ['Create events (draft)',   'Y','Y','Y','Y','-','-','-','-','-'],
        ['Edit any event',          'Y','Y','Y','-','-','-','-','-','-'],
        ['Edit own events',         'Y','Y','Y','Y','-','-','-','-','-'],
        ['Publish / Deactivate',    'Y','Y','Y','-','-','-','-','-','-'],
        ['Delete events',           'Y','Y','-','-','-','-','-','-','-'],
        ['Read all profiles',       'Y','Y','-','-','-','-','-','-','-'],
        ['Update any profile role', 'Y','Y*','-','-','-','-','-','-','-'],
        ['Delete profiles',         'Y','-','-','-','-','-','-','-','-'],
    ]

    table = doc.add_table(rows=1 + len(features), cols=1 + len(roles))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    hdr.cells[0].text = 'Feature'
    set_cell_shading(hdr.cells[0], '1A365D')
    hdr.cells[0].paragraphs[0].runs[0].font.color.rgb = WHITE if hdr.cells[0].paragraphs[0].runs else WHITE
    for p in hdr.cells[0].paragraphs:
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.font.size = Pt(7)
            r.bold = True

    for i, role in enumerate(roles):
        cell = hdr.cells[i + 1]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(role)
        run.bold = True
        run.font.size = Pt(7)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A365D')

    for r_idx, row_data in enumerate(features):
        feature = row_data[0]
        vals = row_data[1:]
        cells = table.rows[r_idx + 1].cells
        cells[0].text = ''
        p = cells[0].paragraphs[0]
        run = p.add_run(feature)
        run.font.size = Pt(7)

        for c_idx, val in enumerate(vals):
            cell = cells[c_idx + 1]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.bold = True
            if val == 'Y':
                run.font.color.rgb = GREEN
            elif val == '-':
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            elif val == 'Y*':
                run.font.color.rgb = RGBColor(0xEA, 0xB3, 0x08)

        if r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, 'F0F4F8')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Y = Allowed  |  Y* = Allowed except cannot promote to GlobalAdmin  |  - = Not allowed')
    run.font.size = Pt(8)
    run.font.color.rgb = MED_GRAY

    doc.add_page_break()

    # ===== 14. PERFORMANCE =====
    doc.add_heading('14. Performance & Accessibility', level=1)

    doc.add_heading('Performance Optimizations', level=2)
    add_styled_table(doc,
        ['Technique', 'Implementation'],
        [
            ['Image optimization', 'WebP with JPEG fallback via <picture>'],
            ['Responsive images', 'Multiple sizes (800w, 1600w)'],
            ['Lazy loading', 'loading="lazy" on below-fold images'],
            ['CSS/JS minification', 'make minify strips comments and whitespace'],
            ['Passive listeners', 'Scroll handlers use { passive: true }'],
            ['Deferred scripts', 'All <script> tags use defer attribute'],
            ['JSON fallback', 'Prevents blank page if Supabase is slow'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('Accessibility Features', level=2)
    add_styled_table(doc,
        ['Feature', 'Implementation'],
        [
            ['Skip navigation', 'Hidden link, visible on focus'],
            ['Semantic HTML', '<header>, <nav>, <main>, <section>, <footer>'],
            ['ARIA attributes', 'aria-label, aria-expanded, aria-modal, aria-live'],
            ['Keyboard navigation', 'Escape closes modals, tab order preserved'],
            ['Color contrast', 'Gold on dark passes WCAG AA'],
            ['Focus indicators', 'outline: 2px solid var(--accent) on focus'],
            ['Reduced motion', 'prefers-reduced-motion media query'],
        ],
        col_widths=[5, 12]
    )

    doc.add_page_break()

    # ===== 15. FUTURE =====
    doc.add_heading('15. Future Considerations', level=1)

    doc.add_heading('Potential Enhancements', level=2)
    enhancements = [
        'Event Registration \u2014 RSVP/ticket system with attendee tracking',
        'Member Directory \u2014 Searchable member list (Volunteer+ access)',
        'Donation Tracking \u2014 Integration with payment gateway (Stripe/PayPal)',
        'Gallery Management \u2014 Admin UI for uploading/organizing photos',
        'Vendor Management \u2014 Admin UI for vendor CRUD operations',
        'Email Notifications \u2014 Event reminders, registration confirmations',
        'PWA Support \u2014 Service worker for offline access',
        'Analytics Dashboard \u2014 Event attendance, member growth metrics',
    ]
    for e in enhancements:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('Scaling Thresholds', level=2)
    add_styled_table(doc,
        ['Threshold', 'Action'],
        [
            ['> 500 monthly active users', 'Monitor Supabase free tier limits (50K auth, 500MB DB)'],
            ['> 10,000 page views/month', 'Consider CDN for static assets'],
            ['> 100 events', 'Add pagination to admin event list'],
            ['Need server-side logic', 'Supabase Edge Functions (Deno)'],
            ['Need real-time features', 'Supabase Realtime subscriptions'],
        ],
        col_widths=[5.5, 11.5]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'This document should be reviewed and updated when significant architectural changes '
        'are made to the system.'
    )
    run.italic = True
    run.font.color.rgb = MED_GRAY

    # ===== SAVE =====
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, '..', 'docs', 'AIKYAM_Architecture.docx')
    out_path = os.path.abspath(out_path)
    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    main()
